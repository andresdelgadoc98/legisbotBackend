from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear un nuevo documento
doc = Document()

# Configurar márgenes (opcional)
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Agregar la fecha (alineada a la derecha)
fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fecha.add_run('Ciudad de México, 25 de octubre de 2023\n\n').bold = True

# Agregar el destinatario (alineado a la izquierda)
destinatario = doc.add_paragraph()
destinatario.add_run('Lic. Juan Pérez\n').bold = True
destinatario.add_run('Director General\n')
destinatario.add_run('Empresa Ejemplo, S.A. de C.V.\n')
destinatario.add_run('Av. Reforma #123, Col. Centro\n')
destinatario.add_run('Ciudad de México, C.P. 06000\n\n')

# Agregar el saludo inicial
doc.add_paragraph('Estimado Lic. Pérez:\n')

# Agregar el cuerpo de la carta
cuerpo = doc.add_paragraph()
cuerpo.add_run('Por medio de la presente, me dirijo a usted para ')
cuerpo.add_run('solicitar información detallada').bold = True
cuerpo.add_run(' sobre los servicios que ofrece su empresa, ')
cuerpo.add_run('específicamente en el área de consultoría financiera.').bold = True
cuerpo.add_run('\n\nAgradezco de antemano su atención y quedo a la espera de su pronta respuesta.\n\n')

# Agregar la despedida
doc.add_paragraph('Atentamente,\n\n')

# Agregar la firma (alineada a la izquierda)
firma = doc.add_paragraph()
firma.add_run('________________________\n').bold = True
firma.add_run('Lic. Ana López\n').bold = True
firma.add_run('Gerente de Proyectos\n')
firma.add_run('Empresa Soluciones Integrales\n')
firma.add_run('Tel: 55-1234-5678\n')
firma.add_run('Email: ana.lopez@soluciones.com\n')

# Guardar el documento
doc.save('carta_formal.docx')

print("Carta generada con éxito.")