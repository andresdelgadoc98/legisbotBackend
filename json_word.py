from docx import Document

def generar_documento_rescisión(nombre_trabajador, fechas_inasistencia, fecha_pago,fecha_actual):

    documento = Document()

    documento.add_heading('Notificación de Rescisión de Contrato Individual de Trabajo', level=1)

    cuerpo = f"""
    Por medio de la presente, le notificamos que con esta fecha, {fecha_actual}, se ha decidido rescindir su contrato individual de trabajo que tenía con esta empresa, con fundamento en la fracción X del artículo 47 de la Ley Federal del Trabajo.

    Lo anterior, debido a que usted ha incurrido en más de tres faltas de asistencia injustificadas en un período de treinta días, específicamente los días {', '.join(fechas_inasistencia)}.

    Por lo tanto, y con base en lo estipulado en el artículo 47 de la Ley Federal del Trabajo, se ha procedido a la rescisión de su contrato individual de trabajo, sin responsabilidad para el patrón.

    Se le informa que su finiquito, que incluye el pago de los días trabajados, la parte proporcional del aguinaldo, vacaciones y prima vacacional, estará a su disposición en las oficinas de la empresa a partir del día {fecha_pago}.
    """

    documento.add_paragraph(cuerpo)

    pie_pagina = documento.sections[0].footer
    parrafo_pie_pagina = pie_pagina.paragraphs[0]
    parrafo_pie_pagina.text = "Atentamente,\n[Nombre y Firma del Representante Legal del Patrón]\n[Cargo del Representante Legal]\n[Nombre de la Empresa]"
    documento.save(f'rescisión_{nombre_trabajador.replace(" ", "_")}.docx')

nombre_trabajador = "Juan Pérez"
fechas_inasistencia = ["01/01/2024", "05/01/2024", "10/01/2024"]
fecha_pago = "15/01/2024"

generar_documento_rescisión(nombre_trabajador, fechas_inasistencia, fecha_pago,fecha_pago)